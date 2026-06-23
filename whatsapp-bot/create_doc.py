from docx import Document

def create_doc():
    doc = Document()
    doc.add_heading('WhatsApp Chatbot Project Documentation', 0)
    
    doc.add_heading('1. Overview', level=1)
    doc.add_paragraph(
        "We have developed a production-ready backend for a WhatsApp sales chatbot. "
        "The system handles user incoming messages, detects intents using AI, provides automated "
        "catalog browsing, checks prices, generates professional quotations in PDF format, and "
        "can escalate conversations to human agents."
    )

    doc.add_heading('2. What We Have Done So Far', level=1)
    doc.add_paragraph("The key functionalities implemented include:", style='List Bullet')
    
    ul1 = doc.add_paragraph("Modular Project Structure: Configured a robust backend using FastAPI.", style='List Bullet')
    ul2 = doc.add_paragraph("Intelligent Intent Classification: Integrated Anthropic's Claude AI to dynamically process and classify incoming user queries (e.g., browse_categories, request_quote, talk_to_human, etc.).", style='List Bullet')
    ul3 = doc.add_paragraph("Session Management: Handled persistent state using a SQLite database.", style='List Bullet')
    ul4 = doc.add_paragraph("Catalog & Quoting Engine: Set up functionalities for users to view product catalogs and built a service to generate PDF quotes directly within WhatsApp using ReportLab.", style='List Bullet')
    ul5 = doc.add_paragraph("Human Handoff Protocol: Designed logic to smoothly bypass the bot and assign a human agent when requested or when an intent is unrecognized.", style='List Bullet')

    doc.add_heading('3. Tech Stack', level=1)
    doc.add_paragraph("Core Languages & Frameworks:", style='List Bullet')
    doc.add_paragraph("Python", style='List Bullet 2')
    doc.add_paragraph("FastAPI (Backend routing and REST architecture)", style='List Bullet 2')

    doc.add_paragraph("AI Services:", style='List Bullet')
    doc.add_paragraph("Anthropic Claude API (For Natural Language Processing & Intent Classification)", style='List Bullet 2')

    doc.add_paragraph("Data & Storage:", style='List Bullet')
    doc.add_paragraph("SQLite (Database backend for states & simple storage)", style='List Bullet 2')
    doc.add_paragraph("SQLAlchemy (ORM)", style='List Bullet 2')

    doc.add_paragraph("Other Libraries:", style='List Bullet')
    doc.add_paragraph("ReportLab (For generating PDF quotations dynamically)", style='List Bullet 2')
    doc.add_paragraph("Uvicorn (ASGI web server)", style='List Bullet 2')

    doc.add_heading('4. Tools Used', level=1)
    doc.add_paragraph("Visual Studio Code (or standard Python IDEs)", style='List Bullet')
    doc.add_paragraph("Git / Version Control", style='List Bullet')
    doc.add_paragraph("WhatsApp Business API (or proxy via Meta for Webhooks/Messaging)", style='List Bullet')
    doc.add_paragraph("Postman / cURL for endpoint testing", style='List Bullet')

    doc.save('WhatsApp_Chatbot_Documentation.docx')
    print("Document successfully created: WhatsApp_Chatbot_Documentation.docx")

if __name__ == '__main__':
    create_doc()
